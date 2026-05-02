import json
from datetime import date
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = PROJECT_ROOT / "outputs" / "reports"
FIGURE_DIR = PROJECT_ROOT / "outputs" / "figures"
DOCX_OUT = REPORT_DIR / "Product_Visual_Search_Full_Report.docx"
MD_OUT = REPORT_DIR / "Product_Visual_Search_Full_Report.md"
SUMMARY_CSV = REPORT_DIR / "experiment_summary.csv"
TITLE = "Product Visual Search and Retrieval Using ResNet18 and CLIP"


def load_json(path):
    path = Path(path)
    if not path.exists():
        print(f"missing_json={path}")
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value, digits=4):
    try:
        if pd.isna(value):
            return "N/A"
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def get_font(size=28, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    words = str(text).split()
    lines, current = [], ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def create_table_image(headers, rows, out_path, title=None):
    col_widths = [260] + [180] * (len(headers) - 1)
    if len(headers) >= 9:
        col_widths = [280, 190, 180, 180, 160, 160, 160, 160, 170]
    row_height = 82
    title_height = 52 if title else 0
    margin = 24
    width = sum(col_widths) + margin * 2
    height = title_height + row_height * (len(rows) + 1) + margin * 2
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(32, True)
    header_font = get_font(25, True)
    body_font = get_font(25, False)
    y = margin
    if title:
        draw.text((margin, y), title, fill=(31, 78, 121), font=title_font)
        y += title_height
    for r_idx, values in enumerate([headers] + rows):
        x = margin
        for c_idx, value in enumerate(values):
            fill = (31, 78, 121) if r_idx == 0 else (248, 250, 252)
            text_color = "white" if r_idx == 0 else (30, 30, 30)
            font = header_font if r_idx == 0 else body_font
            draw.rectangle([x, y, x + col_widths[c_idx], y + row_height], fill=fill, outline=(180, 190, 200), width=2)
            lines = wrap_text(draw, value, font, col_widths[c_idx] - 18)
            line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + 4
            text_y = y + max(4, (row_height - line_height * len(lines)) // 2)
            for line in lines:
                line_width = draw.textbbox((0, 0), line, font=font)[2]
                draw.text((x + (col_widths[c_idx] - line_width) / 2, text_y), line, fill=text_color, font=font)
                text_y += line_height
            x += col_widths[c_idx]
        y += row_height
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path)
    return out_path


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(64, 64, 64)


def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_image(doc, path, caption, width=6.2):
    path = Path(path)
    if not path.exists():
        print(f"skip_missing_figure={path}")
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    return True


def gather_results():
    summary = pd.read_csv(SUMMARY_CSV)
    full_stats = pd.read_csv(REPORT_DIR / "full_dataset_stats.csv")
    train_log = pd.read_csv(REPORT_DIR / "resnet18_full_training_log.csv")
    train_log = train_log[train_log.get("is_smoke_test", False).astype(str).str.lower() != "true"].copy()
    test_metrics = load_json(REPORT_DIR / "resnet18_full_test_metrics.json")
    retrieval_metrics = load_json(REPORT_DIR / "resnet18_full_retrieval_metrics.json")
    dropped = pd.read_csv(REPORT_DIR / "dropped_rare_classes.csv")
    best_row = train_log.loc[train_log["val_top1"].idxmax()] if not train_log.empty else None
    return summary, full_stats, train_log, test_metrics, retrieval_metrics, dropped, best_row


def build_markdown(summary, full_stats, train_log, test_metrics, retrieval_metrics, dropped, best_row):
    comparison = []
    for model_name, label in [
        ("resnet18", "ResNet18 debug"),
        ("clip_vit_b_32_openai_local", "CLIP debug"),
        ("clip_vit_b_32_openai_full", "Frozen CLIP full dataset"),
        ("clip_vit_b_32_articletype_lite", "Fine-tuned CLIP full dataset"),
        ("resnet18_full_dataset", "ResNet18 full dataset"),
    ]:
        rows = summary[summary["model_name"] == model_name]
        if rows.empty:
            continue
        row = rows.iloc[-1]
        comparison.append(
            f"| {label} | {row.get('num_classes', 'N/A')} | {row.get('train_size', 'N/A')} | {row.get('test_size', 'N/A')} | "
            f"{fmt(row.get('top1_acc'))} | {fmt(row.get('top5_acc'))} | {fmt(row.get('recall@1'))} | {fmt(row.get('recall@5'))} | {fmt(row.get('recall@10'))} |"
        )
    return f"""# {TITLE}

Course Project Report  
Date: {date.today().strftime('%B %d, %Y')}  
Author: Student Name

## Abstract

This project implements an end-to-end product visual search pipeline. It includes a supervised ResNet18 baseline, a frozen CLIP ViT-B/32 retrieval model, and a lightweight fine-tuned CLIP ViT-B/32 model trained with articleType supervision. Retrieval is evaluated with Recall@K because product search returns ranked candidate products rather than only one class label.

The full ResNet18 model reaches test Top-1 accuracy of {fmt(test_metrics.get('test_top1_acc'))} and test Top-5 accuracy of {fmt(test_metrics.get('test_top5_acc'))}. In full retrieval, ResNet18 achieves Recall@1 / Recall@5 / Recall@10 of {fmt(retrieval_metrics.get('recall@1'))} / {fmt(retrieval_metrics.get('recall@5'))} / {fmt(retrieval_metrics.get('recall@10'))}. Frozen CLIP and fine-tuned CLIP use the same full train/test split, allowing fair comparison without query/gallery mismatch.

## Introduction

Product visual search allows users to find visually similar products from images instead of keywords. This is important for e-commerce, social shopping, and content platforms because users often know what a product looks like before they know its exact name. A useful system should return a ranked set of visually similar products, not only a category label.

This project therefore combines classification and retrieval. Classification validates whether a model understands product categories. Retrieval tests whether the learned representation can support search and recommendation. The main technical idea is to convert every product image into an embedding vector, compare embeddings with cosine similarity, and return the Top-K nearest gallery products.

## Dataset and Preprocessing

- Valid full dataset images: {test_metrics.get('train_size', 0) + test_metrics.get('val_size', 0) + test_metrics.get('test_size', 0)}
- Number of articleType classes: {len(full_stats)}
- Train / val / test: {test_metrics.get('train_size', 'N/A')} / {test_metrics.get('val_size', 'N/A')} / {test_metrics.get('test_size', 'N/A')}
- Dropped rare classes: {len(dropped)}
- Main image source: `data/raw/images/`

The project avoids the duplicate extracted folder under `data/raw/myntradataset/images/` to reduce train/test leakage. Rare classes with fewer than three valid images are removed before splitting. Train images are used as the gallery; test images are used as query images.

## Methodology

ResNet18 is used as a supervised baseline trained on articleType labels. Its final classification layer is replaced for the dataset classes, and the penultimate 512-dimensional feature vector is used as the retrieval embedding.

CLIP is evaluated in two ways. Frozen CLIP uses pretrained visual-semantic embeddings without local training. Fine-tuned CLIP starts from the same pretrained weights but adapts part of the image encoder with articleType supervision.

All model routes convert images into embeddings and use cosine similarity for Top-K product retrieval. Recall@K measures whether a same-articleType item appears in the first K retrieved products.

## Full ResNet18 Results

- Best epoch: {int(best_row['epoch']) if best_row is not None else 'N/A'}
- Best val Top-1: {fmt(best_row['val_top1'] if best_row is not None else None)}
- Best val Top-5: {fmt(best_row['val_top5'] if best_row is not None else None)}
- Test loss: {fmt(test_metrics.get('test_loss'))}
- Test Top-1: {fmt(test_metrics.get('test_top1_acc'))}
- Test Top-5: {fmt(test_metrics.get('test_top5_acc'))}
- Full Recall@1: {fmt(retrieval_metrics.get('recall@1'))}
- Full Recall@5: {fmt(retrieval_metrics.get('recall@5'))}
- Full Recall@10: {fmt(retrieval_metrics.get('recall@10'))}
- Full gallery size: {retrieval_metrics.get('gallery_size', 'N/A')}

## Comparison

| Model | Classes | Train/Gallery | Test/Query | Top-1 | Top-5 | Recall@1 | Recall@5 | Recall@10 |
|---|---:|---:|---:|---:|---:|---:|---:|---:|
{chr(10).join(comparison)}

## Product Demo

The Streamlit demo supports image upload, random query selection, Top-K retrieval, and model switching between ResNet18 Full Dataset, Frozen CLIP ViT-B/32, and Fine-tuned CLIP ViT-B/32. The models use the same full split in the app: test images are sampled as queries and train images form the gallery.

This synchronized split is important for a fair demo. Previously, a random query could be a test image for one model but a gallery image for another model. The current design avoids that confusion and makes model switching easier to explain.

## Discussion

ResNet18 is strong at dataset-specific articleType recognition because it is trained directly on the target labels. Frozen CLIP is strong at Top-K recall because its pretrained representation captures broader visual-semantic relationships. Fine-tuned CLIP tests whether lightweight domain adaptation can improve the pretrained model's first-rank retrieval behavior.

## Business Interpretation

The system can support image-based product discovery, similar item recommendation, merchant listing support, and duplicate detection. In a production system, the embedding model would likely act as the first-stage retrieval model, while a ranking layer would combine visual similarity with product metadata such as price, brand, color, inventory, and user preferences.

## Limitations and Future Work

The current relevance metric is articleType-level rather than exact SKU-level similarity. The system does not yet use FAISS, metadata ranking, or a production backend. Future work should add approximate nearest-neighbor indexing, stronger metric-learning objectives, product metadata ranking, SKU-level retrieval evaluation when labels are available, and a production web service.
"""


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    summary, full_stats, train_log, test_metrics, retrieval_metrics, dropped, best_row = gather_results()

    comparison_headers = ["Model", "Classes", "Train", "Test", "Top-1", "Top-5", "R@1", "R@5", "R@10"]
    comparison_rows = []
    for model_name, label in [
        ("resnet18", "ResNet18 debug"),
        ("clip_vit_b_32_openai_local", "CLIP debug"),
        ("clip_vit_b_32_openai_full", "Frozen CLIP full"),
        ("clip_vit_b_32_articletype_lite", "Fine-tuned CLIP full"),
        ("resnet18_full_dataset", "ResNet18 full"),
    ]:
        rows = summary[summary["model_name"] == model_name]
        if not rows.empty:
            row = rows.iloc[-1]
            comparison_rows.append([
                label,
                row.get("num_classes", "N/A"),
                row.get("train_size", "N/A"),
                row.get("test_size", "N/A"),
                fmt(row.get("top1_acc")),
                fmt(row.get("top5_acc")),
                fmt(row.get("recall@1")),
                fmt(row.get("recall@5")),
                fmt(row.get("recall@10")),
            ])
    table_img = create_table_image(
        comparison_headers,
        comparison_rows,
        REPORT_DIR / "full_report_comparison_table.png",
        title="Model Comparison",
    )

    md = build_markdown(summary, full_stats, train_log, test_metrics, retrieval_metrics, dropped, best_row)
    MD_OUT.write_text(md, encoding="utf-8")

    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(31, 78, 121)
    for line in ["Course Project Report", f"Date: {date.today().strftime('%B %d, %Y')}", "Author: Student Name"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(13)
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_p(doc, "This project implements an end-to-end product visual search and retrieval pipeline for fashion product images. The system compares supervised convolutional representation learning with ResNet18, frozen pretrained visual-semantic retrieval with CLIP ViT-B/32, and lightweight supervised CLIP fine-tuning. Instead of treating the task as ordinary image classification, the project frames product understanding as a search problem: a query image is embedded into a vector representation and compared against a gallery of product embeddings to retrieve visually similar items.")
    add_p(doc, "The final pipeline uses the full cleaned dataset of 44,405 valid images across 132 articleType categories. ResNet18, frozen CLIP, and fine-tuned CLIP are evaluated on the same full train/test split. Retrieval quality is measured with Recall@K, which better reflects real product search behavior than classification accuracy alone.")
    add_p(doc, "On the full dataset, ResNet18 reaches a test Top-1 accuracy of 0.8947 and Recall@1 of 0.9016. Frozen CLIP achieves Recall@5 of 0.9543 and Recall@10 of 0.9726, the strongest broad top-k recall. Fine-tuned CLIP reaches test Top-1 accuracy of 0.8667 and improves Recall@1 over frozen CLIP from 0.8289 to 0.8613, showing the effect of lightweight domain adaptation.")

    doc.add_heading("Introduction", level=1)
    add_p(doc, "Product visual search is a common capability in modern e-commerce and content platforms. A user may see a product in a social media post, a short video, a screenshot, or a catalog image, but may not know the exact brand, product title, or keywords needed to search for it. In these cases, image-based search provides a more natural interaction: the user supplies an image, and the system returns a ranked list of visually similar products.")
    add_p(doc, "This task is related to image classification, but it is not the same problem. Classification asks the model to assign a single label such as 'Tshirts' or 'Flip Flops'. Visual search asks the system to retrieve multiple relevant products from a gallery. A perfect class prediction does not guarantee useful search results, and a visually useful search result may still have small differences in color, shape, brand, or style. Therefore, the project evaluates both supervised classification metrics and retrieval-oriented Recall@K metrics.")
    add_p(doc, "The motivation of this project is to connect convolutional neural networks, representation learning, pretrained foundation models, and search retrieval in a realistic product discovery setting. ResNet18 is used as a supervised baseline that learns dataset-specific articleType boundaries. Frozen CLIP is used as a pretrained foundation model, while fine-tuned CLIP adapts the pretrained representation to the fashion product domain.")
    add_p(doc, "The final deliverable is not only a notebook experiment. The project also includes reusable scripts for full-dataset training, checkpointing, evaluation, and index construction, plus a Streamlit demo that allows users to upload a product image or choose a random query image and compare CLIP and ResNet18 retrieval results interactively.")
    doc.add_page_break()

    doc.add_heading("Dataset and Preprocessing", level=1)
    add_p(doc, "The project uses the Fashion Product Images dataset stored locally under the project data directory. The primary image source is `data/raw/images/`, and the metadata file is `data/raw/styles.csv`. The label used throughout the project is `articleType`, which provides product category labels such as Tshirts, Shirts, Casual Shoes, Watches, Flip Flops, Backpacks, and many others.")
    add_p(doc, "The raw extraction contains a duplicate-style path under `data/raw/myntradataset/images/`. The project intentionally avoids using that duplicate folder as the canonical source. This choice matters because duplicate image paths can cause train/test leakage: if the same image appears in both the gallery and query sets, retrieval metrics become artificially high. The final split and indexes use `data/raw/images/` as the canonical source.")
    add_bullets(doc, [
        f"Valid full dataset images: {test_metrics.get('train_size', 0) + test_metrics.get('val_size', 0) + test_metrics.get('test_size', 0)}",
        f"Number of articleType classes: {len(full_stats)}",
        f"Train / validation / test sizes: {test_metrics.get('train_size', 'N/A')} / {test_metrics.get('val_size', 'N/A')} / {test_metrics.get('test_size', 'N/A')}",
        f"Dropped rare classes: {len(dropped)}",
        "The canonical source is data/raw/images/; duplicate extracted image folders are avoided to reduce train/test leakage.",
    ])
    add_p(doc, "Rare articleType categories with fewer than three valid images are removed before splitting because they cannot support a reliable stratified train/validation/test split. After filtering, 132 articleType classes remain. The split strategy assigns approximately 70% of each class to training, 15% to validation, and 15% to testing. The train split acts as the retrieval gallery, while the test split acts as the query set.")
    add_p(doc, "All models use 224 by 224 resized RGB images. ResNet18 uses ImageNet mean and standard deviation normalization. CLIP uses the preprocessing transform returned by OpenCLIP for the ViT-B/32 model. These preprocessing steps are necessary because both models expect images to be presented in the same distribution as their training or pretraining environment.")

    doc.add_heading("Methodology", level=1)
    doc.add_heading("ResNet18 Supervised Baseline", level=2)
    add_p(doc, "ResNet18 is a residual convolutional neural network. Its residual connections allow layers to learn corrections to earlier representations, which makes optimization easier than a plain deep convolutional network. In this project, ResNet18 is initialized with ImageNet pretrained weights, then its final fully connected layer is replaced with a new classification head matching the number of articleType classes.")
    add_p(doc, "The supervised training objective is cross-entropy loss. During training, the model learns to classify product images into articleType categories. After training, the final classification layer is not used for retrieval. Instead, the 512-dimensional feature vector before the final classifier is extracted and L2-normalized. This vector becomes the image embedding used for similarity search.")
    doc.add_heading("CLIP-based Retrieval and Fine-tuning", level=2)
    add_p(doc, "CLIP, or Contrastive Language-Image Pre-training, learns image and text representations from large-scale image-text pairs. The frozen CLIP route uses CLIP ViT-B/32 through OpenCLIP, loaded from a local checkpoint, without training on the current product dataset. This tests whether a pretrained vision-language model can serve as a strong retrieval embedding model before task-specific adaptation.")
    add_p(doc, "The project also includes a lightweight fine-tuned CLIP route. It starts from the same pretrained CLIP weights, freezes most parameters, and trains a small trainable portion plus a classification head using articleType labels. This tests whether supervised adaptation can improve first-rank retrieval while keeping training cost manageable.")
    doc.add_heading("Image Embedding and Cosine Similarity Retrieval", level=2)
    add_p(doc, "For retrieval, every gallery image is converted into an embedding vector and stored in an index file. At search time, the query image is encoded with the selected model, normalized, and compared against gallery embeddings using cosine similarity. The Top-K most similar gallery items are returned to the user.")
    add_p(doc, "Recall@K measures whether at least one item with the same articleType appears in the top K retrieved results. Recall@1 is strict because only the first retrieved item counts. Recall@5 and Recall@10 better reflect search and recommendation scenarios, where users inspect a set of candidate products rather than only one result.")
    doc.add_page_break()

    doc.add_heading("System Implementation", level=1)
    add_p(doc, "The implementation is organized around a report notebook plus reusable scripts. The notebook remains a course project artifact that explains the pipeline, shows outputs, and summarizes results. Long-running full-dataset work is deliberately moved into scripts so that training can be resumed safely and does not require running the entire notebook.")
    add_bullets(doc, [
        "`scripts/prepare_full_dataset_split.py` prepares the full cleaned split.",
        "`scripts/train_resnet18_full.py` trains ResNet18 with checkpoint and resume support.",
        "`scripts/evaluate_resnet18_full.py` evaluates classification performance on the test split.",
        "`scripts/build_resnet18_full_index.py` builds the ResNet18 train-gallery embedding index.",
        "`scripts/evaluate_resnet18_full_retrieval.py` computes full retrieval Recall@K.",
        "`scripts/build_clip_full_index.py` builds the CLIP train-gallery embedding index on the same split.",
        "`scripts/evaluate_clip_full_retrieval.py` evaluates CLIP retrieval on the same full query set.",
        "`scripts/train_clip_articletype.py` fine-tunes CLIP with articleType supervision.",
        "`scripts/evaluate_clip_articletype.py` evaluates the fine-tuned CLIP classifier.",
        "`scripts/build_clip_articletype_index.py` builds the fine-tuned CLIP train-gallery embedding index.",
        "`scripts/evaluate_clip_articletype_retrieval.py` evaluates fine-tuned CLIP retrieval.",
        "`app.py` provides the Streamlit product demo with model switching.",
    ])
    add_p(doc, "The Streamlit application uses cached gallery indexes instead of computing gallery embeddings at request time. This makes the app responsive and avoids accidental long-running computation during a product demo. The app also validates whether indexes are smoke-test artifacts or full indexes, which reduces the risk of presenting incomplete results as final.")

    doc.add_heading("Experimental Setup", level=1)
    add_bullets(doc, [
        "ResNet18 debug: 20 classes, 50 images per class.",
        "CLIP debug: frozen ViT-B/32 image encoder on the debug subset.",
        "ResNet18 full: 132 articleType classes, ImageNet initialization, AdamW, 20 epochs, CUDA mixed precision.",
        "CLIP full: frozen ViT-B/32 image encoder using the same full train-gallery and test-query split as ResNet18 full.",
        "Fine-tuned CLIP full: OpenAI CLIP ViT-B/32 initialization, articleType supervision, lightweight trainable parameter ratio.",
    ])
    add_p(doc, "The full ResNet18 model is trained for 20 epochs with batch size 64. CUDA is used when available, and mixed precision is enabled when supported. The training script saves both a last checkpoint and the best validation Top-1 checkpoint.")
    add_p(doc, "Frozen CLIP does not train any model. Fine-tuned CLIP trains a small part of the pretrained CLIP image encoder plus a classification head, then uses the adapted embedding for retrieval.")

    doc.add_heading("Results", level=1)
    add_image(doc, table_img, "Table 1. Comparison of debug and full-dataset experiments.", width=6.5)
    add_image(doc, FIGURE_DIR / "model_recall_comparison.png", "Figure 1. Original debug Recall@K comparison.")
    add_image(doc, FIGURE_DIR / "model_recall_comparison_full.png", "Figure 2. Recall@K comparison including ResNet18 full dataset.")
    add_p(doc, "The comparison table shows that ResNet18 has the strongest Recall@1 among the full-scope models, while frozen CLIP achieves stronger Recall@5 and Recall@10. Fine-tuned CLIP improves Recall@1 over frozen CLIP, showing that domain adaptation helps first-rank retrieval, but it slightly reduces the broader top-k recall of frozen CLIP.")
    add_image(doc, FIGURE_DIR / "resnet18_full_training_curves.png", "Figure 3. ResNet18 full training curves.")
    add_p(doc, "The training curves show rapid improvement during early epochs followed by more gradual validation gains. The best validation Top-1 accuracy occurs at epoch 10, while the best validation Top-5 accuracy occurs earlier. This pattern is reasonable for a supervised product classifier: the model quickly learns broad category structure and then slowly improves fine category boundaries.")
    add_image(doc, FIGURE_DIR / "resnet18_full_retrieval_summary.png", "Figure 4. ResNet18 full retrieval summary.")
    add_p(doc, "The full ResNet18 retrieval results are computed with train images as the gallery and test images as queries. This avoids evaluation leakage and gives a realistic estimate of whether the model can retrieve similar unseen products from a known catalog.")
    for i, img in enumerate(sorted((FIGURE_DIR / "clip_retrieval_examples").glob("*.png"))[:5], start=1):
        add_image(doc, img, f"Figure {4 + i}. CLIP retrieval example.", width=6.2)
    doc.add_page_break()

    doc.add_heading("Product Demo", level=1)
    add_p(doc, "The Streamlit app turns the experiment into an interactive product demo. Users can upload a product image or select a random sample query from the full test split. The app then retrieves visually similar products from the selected model's train-gallery index. Each result card displays the rank, product image, articleType, cosine similarity score, and image id.")
    add_p(doc, "The app supports switching between ResNet18 Full Dataset, Frozen CLIP ViT-B/32, and Fine-tuned CLIP ViT-B/32. These modes use the same full train/test split: random sample queries come from the full test split, while retrieval searches a full train-gallery index.")
    add_p(doc, "In a real product deployment, this Streamlit app would correspond to the visual search frontend. A production version would likely move embedding extraction and vector search to a backend service, use FAISS or another approximate nearest-neighbor index, and combine visual similarity with business ranking signals such as price, brand, inventory, click-through rate, and personalization.")

    doc.add_heading("Discussion", level=1)
    add_p(doc, "The experiments highlight a useful distinction between supervised category learning and pretrained visual-semantic retrieval. ResNet18 is trained directly on articleType labels, so it learns dataset-specific category boundaries. This helps it achieve strong classification accuracy and strong Recall@1 in the full dataset setting. However, because it is trained to separate articleType classes, it may emphasize category-discriminative features more than broader style or semantic similarity.")
    add_p(doc, "Frozen CLIP is not trained on the local dataset, but it is pretrained on broad image-text pairs. Its full-dataset Recall@5 and Recall@10 are higher than ResNet18 full, suggesting that CLIP is especially strong as a candidate recall model. Fine-tuned CLIP adapts this representation to articleType labels and improves Recall@1 over frozen CLIP, but with a small loss in broader top-k recall.")
    add_p(doc, "The results also show why classification accuracy alone is insufficient. ResNet18 full has strong Top-1 and Top-5 classification accuracy, but the product demo ultimately depends on retrieval behavior. A user does not only care whether the system says 'Flip Flops'; the user wants to see visually similar flip flops, ranked near the top.")

    doc.add_heading("Business Interpretation", level=1)
    add_p(doc, "From a business perspective, this system can be interpreted as the recall layer of a visual product search engine. When a user uploads an image, the system quickly retrieves a set of visually related catalog items. These items could then be re-ranked using structured metadata, user preferences, price, availability, and engagement signals.")
    add_p(doc, "The same embedding infrastructure can support several product features. It can power image-based product discovery, similar item recommendation, automatic product categorization, merchant listing quality checks, and duplicate or near-duplicate detection. On content platforms, it could help connect user-generated images or videos to purchasable products.")
    add_p(doc, "The model comparison suggests a practical hybrid strategy. ResNet18 is useful when the platform has labeled product data and wants a lightweight supervised model. Frozen CLIP is useful when broad generalization and strong Top-K candidate recall are important. Fine-tuned CLIP is useful when the system wants to adapt a pretrained model toward product-category discrimination without training from scratch.")

    doc.add_heading("Limitations", level=1)
    add_bullets(doc, [
        "The relevance signal is articleType-level, not exact SKU-level matching.",
        "The frozen and fine-tuned CLIP results show a tradeoff between broad semantic recall and domain-specific first-rank retrieval.",
        "The system does not yet use FAISS, metadata ranking, or a production backend.",
        "Category imbalance and dataset bias remain important limitations.",
        "The current app is a local Streamlit prototype rather than a deployed service.",
        "The retrieval evaluation does not consider brand, color, gender, season, price, or exact product identity.",
    ])

    doc.add_heading("Future Work", level=1)
    add_bullets(doc, [
        "Scale the full CLIP index with FAISS.",
        "Add FAISS approximate nearest-neighbor search.",
        "Try stronger metric-learning losses for fine-tuned CLIP, such as contrastive loss or triplet loss.",
        "Use product metadata such as brand, color, price, and gender for ranking.",
        "Deploy with FastAPI and React and add a user feedback loop.",
        "Evaluate exact-product or SKU-level retrieval if product identity labels become available.",
        "Add qualitative failure analysis for visually similar but label-mismatched retrievals.",
    ])

    doc.add_heading("Conclusion", level=1)
    add_p(doc, "The project successfully builds a product visual search pipeline from dataset preparation through model training, retrieval evaluation, index construction, and an interactive frontend demo. The final comparison includes ResNet18 supervised training, frozen CLIP retrieval, and lightweight fine-tuned CLIP.")
    add_p(doc, "Overall, the project demonstrates that product recognition and visual retrieval are complementary. Supervised ResNet18 provides strong dataset-specific classification and Recall@1, frozen CLIP provides strong Top-K candidate recall, and fine-tuned CLIP demonstrates how domain adaptation can improve pretrained CLIP's first-rank retrieval behavior.")
    doc.save(DOCX_OUT)
    print(f"docx_created={DOCX_OUT}")
    print(f"markdown_created={MD_OUT}")


if __name__ == "__main__":
    main()
