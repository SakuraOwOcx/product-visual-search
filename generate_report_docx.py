from pathlib import Path
from datetime import date
import os

try:
    import pandas as pd
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from PIL import Image, ImageDraw, ImageFont
except ImportError as exc:
    print("A required reporting dependency is missing.")
    print("Suggested install command:")
    print("pip install python-docx pandas pillow")
    raise


REPORT_DIR = Path("outputs/reports")
FIGURE_DIR = Path("outputs/figures")
SUMMARY_CSV = REPORT_DIR / "experiment_summary.csv"
DOCX_OUT = REPORT_DIR / "Product_Visual_Search_Report.docx"
MD_OUT = REPORT_DIR / "Product_Visual_Search_Report.md"
RECALL_FIGURE = FIGURE_DIR / "model_recall_comparison.png"
CLIP_EXAMPLE_DIR = FIGURE_DIR / "clip_retrieval_examples"
RESULTS_TABLE_IMAGE = REPORT_DIR / "results_table.png"


TITLE = "Product Visual Search and Retrieval Using ResNet18 and CLIP"
TODAY = date.today().strftime("%B %d, %Y")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8.0):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def read_results():
    df = pd.read_csv(SUMMARY_CSV)
    resnet = df[df["model_name"] == "resnet18"].iloc[0].to_dict()
    clip = df[df["model_name"] == "clip_vit_b_32_openai_local"].iloc[0].to_dict()
    return resnet, clip


def format_results_table_rows(resnet, clip):
    return [
        [
            "ResNet18",
            "Supervised training on articleType",
            "512",
            f"{float(resnet['top1_acc']):.4f}",
            f"{float(resnet['top5_acc']):.4f}",
            f"{float(resnet['recall@1']):.4f}",
            f"{float(resnet['recall@5']):.4f}",
            f"{float(resnet['recall@10']):.4f}",
            "18.68s",
        ],
        [
            "CLIP ViT-B/32",
            "Frozen pretrained image encoder",
            "512",
            "N/A",
            "N/A",
            f"{float(clip['recall@1']):.4f}",
            f"{float(clip['recall@5']):.4f}",
            f"{float(clip['recall@10']):.4f}",
            "11.55s",
        ],
    ]


def add_results_table(doc, rows):
    headers = [
        "Model",
        "Training Strategy",
        "Embedding Dim",
        "Top-1 Acc",
        "Top-5 Acc",
        "Recall@1",
        "Recall@5",
        "Recall@10",
        "Runtime",
    ]
    image_path = create_results_table_image(headers, rows)
    add_image_if_exists(
        doc,
        image_path,
        "Table 1. ResNet18 and CLIP debug retrieval results.",
        width=6.4,
    )


def get_font(size=30, bold=False):
    fonts_dir = Path(os.environ.get("WINDIR", "")) / "Fonts"
    font_candidates = [
        fonts_dir / ("calibrib.ttf" if bold else "calibri.ttf"),
        fonts_dir / ("arialbd.ttf" if bold else "arial.ttf"),
    ]
    for font_path in font_candidates:
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size=size)
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    words = str(text).split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def create_results_table_image(headers, rows):
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    col_widths = [230, 455, 175, 165, 165, 165, 165, 175, 160]
    row_heights = [88, 92, 92]
    margin = 24
    width = sum(col_widths) + 2 * margin
    height = sum(row_heights) + 2 * margin

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    header_font = get_font(25, bold=True)
    body_font = get_font(26, bold=False)
    grid_color = (180, 190, 200)
    header_bg = (31, 78, 121)

    y = margin
    for r_idx, values in enumerate([headers] + rows):
        x = margin
        row_height = row_heights[r_idx]
        for c_idx, value in enumerate(values):
            col_width = col_widths[c_idx]
            fill = header_bg if r_idx == 0 else (248, 250, 252)
            draw.rectangle([x, y, x + col_width, y + row_height], fill=fill, outline=grid_color, width=2)

            font = header_font if r_idx == 0 else body_font
            text_color = "white" if r_idx == 0 else (25, 25, 25)
            lines = wrap_text(draw, value, font, col_width - 18)
            line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + 4
            block_height = line_height * len(lines)
            text_y = y + max(6, (row_height - block_height) // 2)
            for line in lines:
                line_width = draw.textbbox((0, 0), line, font=font)[2]
                draw.text((x + (col_width - line_width) / 2, text_y), line, fill=text_color, font=font)
                text_y += line_height
            x += col_width
        y += row_height

    img.save(RESULTS_TABLE_IMAGE)
    return RESULTS_TABLE_IMAGE


def add_image_if_exists(doc, path, caption, width=6.2):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)
        return True
    print(f"Skipping missing figure: {path}")
    return False


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Title"]:
        styles[style_name].font.name = "Calibri"
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.color.rgb = RGBColor(64, 64, 64)


def build_markdown(rows):
    clip_examples = sorted(CLIP_EXAMPLE_DIR.glob("*.png"))[:5]
    table_lines = [
        "| Model | Training Strategy | Embedding Dim | Top-1 Acc | Top-5 Acc | Recall@1 | Recall@5 | Recall@10 | Runtime |",
        "|---|---|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for row in rows:
        table_lines.append("| " + " | ".join(row) + " |")

    md = f"""# {TITLE}

Course Project Report  
Date: {TODAY}  
Author: Student Name

## Abstract

This project builds a product visual search system using the Kaggle Fashion Product Images Small dataset. It compares a supervised ResNet18 baseline with a frozen CLIP ViT-B/32 visual-semantic embedding model. Retrieval is evaluated using Recall@K, which is more aligned with visual search than classification accuracy. On the debug subset, CLIP achieves stronger Recall@5 and Recall@10 than ResNet18.

## Introduction

Product visual search allows users to find visually similar products using an image as the query. E-commerce and content platforms benefit from this capability because users often search by appearance, style, or visual intent rather than exact product names. Classification accuracy alone is insufficient because a visual search system needs to return a ranked list of relevant candidate products. This project extracts image embeddings and retrieves similar products by cosine similarity.

## Dataset and Preprocessing

- Total scanned images: 44,441
- Label source: `articleType` from `styles.csv`
- Debug subset: 20 classes, up to 50 images per class, about 1,000 images
- Train/gallery size: 725
- Validation size: 126
- Test/query size: 149
- Duplicate extracted image folders were handled by prioritizing `data/raw/images/` to avoid train/test leakage.

## Methodology

### ResNet18 Supervised Baseline

ResNet18 is a residual network. Residual connections help train deeper networks by allowing layers to learn corrections to earlier representations. In this project, ResNet18 is trained using `articleType` labels, and the penultimate 512-dimensional feature vector is used as the image embedding for retrieval.

### CLIP-based Retrieval

CLIP stands for Contrastive Language-Image Pre-training. It is a vision-language pretrained model. This project uses only the frozen CLIP image encoder from a local checkpoint: `models/huggingface/timm_vit_base_patch32_clip_224_openai/open_clip_model.safetensors`. CLIP is not fine-tuned on the current dataset.

### Image Retrieval Pipeline

The retrieval pipeline embeds query and gallery images, computes cosine similarity, returns Top-K nearest images, and evaluates Recall@1, Recall@5, and Recall@10.

## Experimental Setup

The experiments use `DEBUG_MODE = True`, 20 classes, at most 50 images per class, a gallery size of 725, and a query size of 149. Both ResNet18 and CLIP produce 512-dimensional embeddings. CUDA/GPU was used. The full 44,441-image dataset was not trained in this report.

## Results

{chr(10).join(table_lines)}

![Recall Comparison](../figures/model_recall_comparison.png)

## Discussion

ResNet18 and CLIP have nearly identical Recall@1, while CLIP has higher Recall@5 and Recall@10. This suggests that CLIP provides a stronger candidate set for Top-K retrieval. ResNet18 is a dataset-specific supervised baseline, while CLIP is a general-purpose pretrained visual-semantic representation. For visual search, Top-K retrieval is more important than a single class prediction because users usually expect a set of similar products.

## Business Interpretation

Product visual search helps users find similar products from images. Higher Recall@K means the system is more likely to return relevant products near the top of the ranking. CLIP's strong performance without training suggests that pretrained foundation models can reduce training cost. ResNet18 remains useful as a lightweight supervised baseline or local customized model.

## Limitations

The experiment uses a debug subset rather than the full dataset. `articleType` is a category-level label, not a fine-grained product identity label. CLIP is not fine-tuned, and the retrieval metric does not consider attributes such as color, brand, style, gender, season, or usage.

## Future Work

Future work includes medium-scale evaluation, full-dataset retrieval, CLIP fine-tuning, product attribute-aware retrieval, FAISS indexing, fine-grained retrieval beyond `articleType`, and deployment as a visual search demo.

## Conclusion

The project successfully constructs an end-to-end product visual search pipeline. ResNet18 validates supervised learning for product recognition and retrieval, while CLIP demonstrates that pretrained vision-language models can provide strong retrieval embeddings without dataset-specific training. CLIP's better Recall@5 and Recall@10 indicates strong practical value for product visual retrieval.

"""
    if clip_examples:
        md += "\n## CLIP Retrieval Examples\n\n"
        for path in clip_examples:
            md += f"![{path.name}](../figures/clip_retrieval_examples/{path.name})\n\n"
    return md


def build_docx():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    resnet, clip = read_results()
    rows = format_results_table_rows(resnet, clip)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    for line in ["Course Project Report", f"Date: {TODAY}", "Author: Student Name"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_paragraph(doc, "This project builds a product visual search system using the Kaggle Fashion Product Images Small dataset. It compares a supervised ResNet18 baseline with a frozen CLIP ViT-B/32 visual-semantic embedding model. Retrieval is evaluated using Recall@K, which is more aligned with visual search than classification accuracy. On the debug subset, CLIP achieves stronger Recall@5 and Recall@10 than ResNet18.")

    doc.add_heading("Introduction", level=1)
    add_paragraph(doc, "Product visual search allows users to find visually similar products using an image as the query. E-commerce and content platforms benefit from this capability because users often search by appearance, style, or visual intent rather than exact product names. Classification accuracy alone is insufficient because a visual search system needs to return a ranked list of relevant candidate products.")

    doc.add_heading("Dataset and Preprocessing", level=1)
    add_bullets(doc, [
        "Total scanned images: 44,441.",
        "Label source: articleType from styles.csv.",
        "Debug subset: 20 classes, up to 50 images per class, about 1,000 images.",
        "Train/gallery size: 725; validation size: 126; test/query size: 149.",
        "Images are resized and normalized before model inference or training.",
        "Duplicate extracted image folders were handled by prioritizing data/raw/images/ to avoid train/test leakage.",
    ])

    doc.add_heading("Methodology", level=1)
    doc.add_heading("ResNet18 Supervised Baseline", level=2)
    add_paragraph(doc, "ResNet18 is a residual network. Residual connections help train deeper networks by allowing layers to learn corrections to earlier representations. The model is trained using articleType labels, then its final classification layer is removed so the penultimate 512-dimensional feature vector can be used as an image embedding.")
    doc.add_heading("CLIP-based Retrieval", level=2)
    add_paragraph(doc, "CLIP stands for Contrastive Language-Image Pre-training. It is a vision-language pretrained model. This project uses only the frozen CLIP image encoder from the local checkpoint models/huggingface/timm_vit_base_patch32_clip_224_openai/open_clip_model.safetensors. CLIP is not fine-tuned on the current dataset.")
    doc.add_heading("Image Retrieval Pipeline", level=2)
    add_paragraph(doc, "The retrieval pipeline embeds query and gallery images, computes cosine similarity, retrieves Top-K nearest images, and evaluates Recall@1, Recall@5, and Recall@10.")

    doc.add_heading("Experimental Setup", level=1)
    add_paragraph(doc, "The experiments use DEBUG_MODE = True, 20 classes, at most 50 images per class, a gallery size of 725, and a query size of 149. Both ResNet18 and CLIP produce 512-dimensional embeddings. CUDA/GPU was used. The full 44,441-image dataset was not trained in this report.")

    doc.add_heading("Results", level=1)
    add_results_table(doc, rows)
    add_image_if_exists(doc, RECALL_FIGURE, "Figure 1. ResNet18 vs CLIP Recall@K comparison.", width=5.8)

    doc.add_heading("CLIP Retrieval Examples", level=2)
    examples = sorted(CLIP_EXAMPLE_DIR.glob("*.png"))[:5]
    if examples:
        for i, path in enumerate(examples, start=1):
            add_image_if_exists(doc, path, f"Figure {i + 1}. CLIP Top-5 retrieval example.", width=6.2)
    else:
        add_paragraph(doc, "No CLIP retrieval example figures were found, so this section is left without images.")

    doc.add_heading("Discussion", level=1)
    add_paragraph(doc, "ResNet18 and CLIP have nearly identical Recall@1, while CLIP has higher Recall@5 and Recall@10. This suggests that CLIP provides a stronger candidate set for Top-K retrieval. ResNet18 is a dataset-specific supervised baseline, while CLIP is a general-purpose pretrained visual-semantic representation. For visual search, Top-K retrieval is more important than a single class prediction because users usually expect a set of similar products.")

    doc.add_heading("Business Interpretation", level=1)
    add_paragraph(doc, "Product visual search helps users find similar products from images. Higher Recall@K means the system is more likely to return relevant products near the top of the ranking. CLIP's strong performance without training suggests that pretrained foundation models can reduce training cost. ResNet18 remains useful as a lightweight supervised baseline or local customized model.")

    doc.add_heading("Limitations", level=1)
    add_bullets(doc, [
        "The experiment uses a debug subset rather than the full dataset.",
        "articleType is a category-level label, not a fine-grained product identity label.",
        "Debug subset results may not fully represent large-scale production behavior.",
        "CLIP is not fine-tuned on the current dataset.",
        "The retrieval metric does not consider color, brand, style, gender, season, or usage.",
    ])

    doc.add_heading("Future Work", level=1)
    add_bullets(doc, [
        "Run medium-scale and full-dataset evaluation.",
        "Fine-tune CLIP on fashion product data.",
        "Use attributes such as color, gender, season, and usage.",
        "Add FAISS for approximate nearest-neighbor search.",
        "Improve fine-grained retrieval beyond articleType.",
        "Deploy the pipeline as a simple visual search demo.",
    ])

    doc.add_heading("Conclusion", level=1)
    add_paragraph(doc, "The project successfully constructs an end-to-end product visual search pipeline. ResNet18 validates supervised learning for product recognition and retrieval, while CLIP demonstrates that pretrained vision-language models can provide strong retrieval embeddings without dataset-specific training. CLIP's better Recall@5 and Recall@10 indicates strong practical value for product visual retrieval.")

    doc.save(DOCX_OUT)
    MD_OUT.write_text(build_markdown(rows), encoding="utf-8")
    print(f"docx_created={DOCX_OUT}")
    print(f"markdown_created={MD_OUT}")
    print(f"clip_examples_inserted={len(examples)}")
    print(f"recall_figure_inserted={RECALL_FIGURE.exists()}")


if __name__ == "__main__":
    build_docx()
